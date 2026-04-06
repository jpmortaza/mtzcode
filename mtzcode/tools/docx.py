"""Tools para ler e escrever documentos .docx (Word)."""
from __future__ import annotations

from pathlib import Path

from pydantic import BaseModel, Field

from mtzcode.tools.base import Tool, ToolError


# ---------- Read ----------


class DocxReadArgs(BaseModel):
    path: str = Field(..., description="Caminho do arquivo .docx a ler.")


class DocxReadTool(Tool):
    name = "docx_read"
    destructive = False
    description = (
        "Lê um arquivo .docx (Word) e devolve o conteúdo em markdown. "
        "Extrai parágrafos e tabelas do documento."
    )
    Args = DocxReadArgs

    def run(self, args: DocxReadArgs) -> str:  # type: ignore[override]
        # Import lazy: só carrega python-docx quando a tool é realmente chamada.
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise ToolError(
                "dependência ausente: instale `python-docx` (pip install python-docx)"
            ) from exc

        p = Path(args.path).expanduser()
        if not p.exists():
            raise ToolError(f"arquivo não encontrado: {p}")

        try:
            doc = Document(str(p))
        except Exception as exc:  # noqa: BLE001
            raise ToolError(f"erro abrindo {p}: {exc}") from exc

        linhas: list[str] = []

        # Parágrafos: tenta inferir heading pelo style.name (ex: "Heading 1").
        for par in doc.paragraphs:
            texto = (par.text or "").rstrip()
            if not texto:
                linhas.append("")
                continue
            style = (par.style.name or "").lower() if par.style else ""
            if "heading 1" in style:
                linhas.append(f"# {texto}")
            elif "heading 2" in style:
                linhas.append(f"## {texto}")
            elif "heading 3" in style:
                linhas.append(f"### {texto}")
            elif "list" in style or "bullet" in style:
                linhas.append(f"- {texto}")
            else:
                linhas.append(texto)

        # Tabelas: cada tabela vira uma tabela markdown simples.
        for i, table in enumerate(doc.tables):
            linhas.append("")
            linhas.append(f"<!-- tabela {i + 1} -->")
            rows = [
                [(cell.text or "").strip().replace("\n", " ") for cell in row.cells]
                for row in table.rows
            ]
            if not rows:
                continue
            header = rows[0]
            linhas.append("| " + " | ".join(header) + " |")
            linhas.append("| " + " | ".join("---" for _ in header) + " |")
            for row in rows[1:]:
                linhas.append("| " + " | ".join(row) + " |")

        return "\n".join(linhas).strip() or "(documento vazio)"


# ---------- Write ----------


class DocxWriteArgs(BaseModel):
    path: str = Field(..., description="Caminho de saída do .docx.")
    content: str = Field(..., description="Conteúdo em markdown simples.")
    title: str | None = Field(None, description="Título opcional do documento (vai como heading 1 no topo).")


class DocxWriteTool(Tool):
    name = "docx_write"
    destructive = True
    description = (
        "Cria um arquivo .docx (Word) a partir de markdown simples. "
        "Suporta `# heading1`, `## heading2`, `- bullet` e parágrafos normais."
    )
    Args = DocxWriteArgs

    def run(self, args: DocxWriteArgs) -> str:  # type: ignore[override]
        # Import lazy.
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise ToolError(
                "dependência ausente: instale `python-docx` (pip install python-docx)"
            ) from exc

        p = Path(args.path).expanduser()
        p.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        if args.title:
            doc.add_heading(args.title, level=1)

        # Parser bem simples linha-a-linha.
        for raw in args.content.splitlines():
            linha = raw.rstrip()
            if not linha.strip():
                # Linha vazia → parágrafo vazio (espaçamento).
                doc.add_paragraph("")
                continue
            if linha.startswith("# "):
                doc.add_heading(linha[2:].strip(), level=1)
            elif linha.startswith("## "):
                doc.add_heading(linha[3:].strip(), level=2)
            elif linha.startswith("### "):
                doc.add_heading(linha[4:].strip(), level=3)
            elif linha.startswith("- "):
                doc.add_paragraph(linha[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(linha)

        try:
            doc.save(str(p))
        except OSError as exc:
            raise ToolError(f"erro salvando {p}: {exc}") from exc

        return f"salvo em {p}"
